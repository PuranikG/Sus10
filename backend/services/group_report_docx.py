"""
Sus10 AI — Group DOCX Report Generator
Produces a Word document (A4, Calibri, no dark backgrounds) for a project portfolio.
Sections: Cover | Executive Summary | Project Overview | Building Table |
          Four Pillars | ESG/BRSR or Community Impact | Methodology | Next Steps
"""

from datetime import date
from typing import Dict, Any, List, Optional
import io


def _fmt_inr(n) -> str:
    """Indian number formatting — e.g. 1234567 → '12,34,567'."""
    n = int(n or 0)
    s = str(n)
    if len(s) <= 3:
        return s
    last3 = s[-3:]
    rest = s[:-3]
    parts = []
    while len(rest) > 2:
        parts.append(rest[-2:])
        rest = rest[:-2]
    if rest:
        parts.append(rest)
    parts.reverse()
    return ",".join(parts) + "," + last3


def generate_group_report_docx(
    group: Dict[str, Any],
    buildings: List[Dict[str, Any]],
    sustenance: Dict[str, Any],
    per_building: List[Dict[str, Any]],
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page setup: A4, 2.5 cm margins ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.5))

    # ── Style helpers ────────────────────────────────────────────────────────
    DARK_GREEN = RGBColor(0x0D, 0x3B, 0x1E)
    BLACK      = RGBColor(0x00, 0x00, 0x00)
    GREY       = RGBColor(0xF3, 0xF4, 0xF6)

    def heading(text: str, level: int = 1) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GREEN
        run.font.size = Pt(14 if level == 1 else 12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)

    def body(text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run.bold = bold
        p.paragraph_format.space_after = Pt(3)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

    def add_table_row(table, cells, bold_first: bool = False, header: bool = False):
        row = table.add_row()
        for i, (cell, text) in enumerate(zip(row.cells, cells)):
            cell.text = str(text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = BLACK
                    if header or (bold_first and i == 0):
                        run.bold = True
            if header:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F3F4F6")
                tcPr.append(shd)

    # ── Derived numbers ──────────────────────────────────────────────────────
    project_name    = group.get("name", "Unnamed Project")
    project_type    = group.get("type", "")
    primary_city    = group.get("primary_city", "")
    today           = date.today().strftime("%d %B %Y")
    bld_count       = len(buildings)
    is_residential  = project_type in ("residential_colony", "apartment_complex")

    s = sustenance if isinstance(sustenance, dict) else {}
    total_solar_kwp   = round(s.get("total_solar_kwp", 0))
    total_kwh         = round(s.get("total_solar_kwh_per_year", 0))
    total_plants      = int(s.get("total_plants_count", 0))
    total_food_kg     = round(s.get("total_food_kg_per_year", 0))
    total_biogas_m3   = round(s.get("total_biogas_m3_per_year", 0))
    total_rain_kl     = round(s.get("total_rainwater_kl_per_year", 0))
    total_rain_l      = total_rain_kl * 1000
    total_savings_inr = round(s.get("total_annual_savings_inr", 0))
    total_co2_kg      = round(s.get("total_co2_offset_kg_per_year", 0))
    total_co2_t       = round(total_co2_kg / 1000, 2)
    total_terrace_sqm = s.get("total_terrace_sqm", 0) or s.get("total_footprint_sqm", 0)
    total_terrace_sqft = round(total_terrace_sqm * 10.764)
    total_footprint_sqm = s.get("total_footprint_sqm", 0)
    brsr_narrative    = s.get("brsr_narrative", "")
    biogas_cylinders  = round(total_biogas_m3 * 0.45 / 14.2) if total_biogas_m3 else 0

    type_label_map = {
        "enterprise": "Enterprise / Brand Chain",
        "developer": "Developer Portfolio",
        "federation": "RWA / Society Federation",
        "rwa": "Single Residential Society",
        "residential_colony": "Residential Colony / Housing Society",
        "apartment_complex": "Apartment Complex / Gated Community",
    }
    type_label = type_label_map.get(project_type, project_type.replace("_", " ").title())

    # ════════════════════════════════════════════════════════════════════
    # SECTION 1 — COVER PAGE
    # ════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f"{project_name}")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_GREEN
    run.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Rooftop Sustainability Assessment")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GREEN

    doc.add_paragraph()
    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_by.add_run("Prepared by Sus10 AI  |  sus10.ai")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK

    p_dt = doc.add_paragraph()
    p_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_dt.add_run(today)
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    doc.add_paragraph()
    p_draft = doc.add_paragraph()
    p_draft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_draft.add_run("DRAFT — For Review and Discussion")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True

    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_foot.add_run("Confidential  |  All figures are estimates based on State 1 (Remote Assessment)")
    run.font.name = "Calibri"
    run.font.size = Pt(9)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 2 — EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════
    heading("2. Executive Summary")
    if brsr_narrative:
        body(brsr_narrative)
    else:
        body(
            f"This assessment covers {bld_count} building{'s' if bld_count != 1 else ''} "
            f"in the {project_name} portfolio ({primary_city}). "
            f"Total rooftop sustainability potential is detailed in the sections below."
        )

    doc.add_paragraph()
    bullet(f"Annual savings potential: Rs.{_fmt_inr(total_savings_inr)}")
    bullet(f"CO2 offset: {_fmt_inr(total_co2_kg)} kg/year")
    bullet(f"{bld_count} building{'s' if bld_count != 1 else ''} assessed under State 1 Remote Estimation")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 3 — PROJECT OVERVIEW
    # ════════════════════════════════════════════════════════════════════
    heading("3. Project Overview")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    overview_rows = [
        ("Project Name",       project_name),
        ("Project Type",       type_label),
        ("Primary City",       primary_city),
        ("Buildings Assessed", str(bld_count)),
        ("Total Terrace Area", f"{_fmt_inr(total_terrace_sqft)} sq ft"),
        ("Assessment Date",    today),
        ("Methodology State",  "State 1 — Remote Estimation"),
    ]
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    for cell in tbl.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)

    for label, value in overview_rows:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 4 — BUILDING-BY-BUILDING SUMMARY TABLE
    # ════════════════════════════════════════════════════════════════════
    heading("4. Building-by-Building Summary")

    cols = ["Building Name", "City", "Terrace (sq ft)", "Solar kWp", "Plants",
            "Biogas (cyl/yr)", "Rainwater (L/yr)", "Annual Savings (Rs.)"]
    tbl2 = doc.add_table(rows=1, cols=len(cols))
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl2.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        cell.text = col
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F3F4F6")
        tcPr.append(shd)

    tot_solar = tot_plants = tot_cyl = tot_rain_l_sum = tot_sav = 0
    for br in per_building:
        bname = br.get("building_name") or "—"
        bcity = br.get("city") or "—"
        bterrace = round((br.get("usable_terrace_sqm") or 0) * 10.764)
        bs = br.get("summary", {})
        bp = br.get("pillars", {})
        bsolar_kwp  = bs.get("solar_kwp", 0)
        bplants     = bs.get("plants_count", 0)
        bgas_m3     = bs.get("biogas_m3_per_year", 0)
        bgas_cyl    = round(float(bgas_m3) * 0.45 / 14.2) if bgas_m3 else 0
        brain_kl    = bs.get("rainwater_kl_per_year", 0)
        brain_l     = round(float(brain_kl) * 1000)
        bsav        = bs.get("total_annual_savings_inr", 0)

        tot_solar  += float(bsolar_kwp)
        tot_plants += int(bplants)
        tot_cyl    += bgas_cyl
        tot_rain_l_sum += brain_l
        tot_sav    += int(bsav)

        row = tbl2.add_row()
        for i, val in enumerate([bname, bcity, _fmt_inr(bterrace), bsolar_kwp,
                                  _fmt_inr(bplants), bgas_cyl,
                                  _fmt_inr(brain_l), _fmt_inr(bsav)]):
            row.cells[i].text = str(val)
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)

    # Totals row
    tot_row = tbl2.add_row()
    totals = ["TOTAL", "", _fmt_inr(total_terrace_sqft), round(tot_solar, 1),
              _fmt_inr(tot_plants), tot_cyl, _fmt_inr(tot_rain_l_sum), _fmt_inr(tot_sav)]
    for i, val in enumerate(totals):
        tot_row.cells[i].text = str(val)
        for para in tot_row.cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(9)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 5 — FOUR PILLARS
    # ════════════════════════════════════════════════════════════════════
    heading("5. Four Sustainability Pillars")

    heading("5.1 Solar PV Generation", level=2)
    bullet(f"Total installed potential: {total_solar_kwp} kWp across {bld_count} building{'s' if bld_count != 1 else ''}")
    bullet(f"Annual generation: {_fmt_inr(total_kwh)} units/year")
    bullet(f"Annual savings: Rs.{_fmt_inr(total_savings_inr)}")
    bullet(f"CO2 offset: {_fmt_inr(total_co2_kg)} kg/year")
    body("Methodology: MNRE Solar Radiation Handbook (GHI), 65% area utilisation, 20% panel efficiency, 0.75 performance ratio, CEA grid emission factor 0.82 kg CO2/kWh.")

    heading("5.2 Rooftop Plantation", level=2)
    bullet(f"Total plants: {_fmt_inr(total_plants)}")
    bullet(f"Annual food yield: {_fmt_inr(total_food_kg)} kg/year")
    bullet(f"CO2 sequestered: {_fmt_inr(s.get('total_co2_offset_kg_per_year', 0))} kg/year (combined solar + plantation)")
    body("Methodology: ICAR KVK container farming guidelines, 1 plant/sq ft at 70% terrace utilisation, 2.5 kg food/plant/year, 3.5 kg CO2 sequestered/plant/year.")

    heading("5.3 Rainwater Harvesting", level=2)
    bullet(f"Annual yield: {_fmt_inr(total_rain_l)} litres/year  ({_fmt_inr(total_rain_kl)} kL/year)")
    bullet(f"Annual savings: Rs.{_fmt_inr(s.get('total_annual_savings_inr', 0))}")
    body("Methodology: IMD long-period rainfall average, CPCB/CGWB runoff coefficient 0.85, 5% first-flush loss.")

    heading("5.4 Biogas from Organic Waste", level=2)
    bullet(f"Annual biogas: {_fmt_inr(total_biogas_m3)} m³/year  (= {_fmt_inr(biogas_cylinders)} LPG cylinder equivalents/year)")
    bullet(f"Annual LPG savings: Rs.{_fmt_inr(s.get('total_annual_savings_inr', 0))}")
    body("Methodology: IS 16190:2014, 0.4 kg organic waste/person/day, 0.08 m³ biogas/kg waste, 1 m³ biogas ≈ 0.45 kg LPG.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 6 — ESG/BRSR (commercial) OR COMMUNITY IMPACT (residential)
    # ════════════════════════════════════════════════════════════════════
    if not is_residential:
        heading("6. ESG Disclosure Mapping — BRSR Principle 6 / GRI Environmental Standards")
        body("All figures are indicative estimates based on remote assessment. Verification requires site audit.", bold=True)
        doc.add_paragraph()

        eri_cols = ["GRI Standard", "Sus10 Estimate", "ESG Relevance"]
        tbl3 = doc.add_table(rows=1, cols=3)
        tbl3.style = "Table Grid"
        tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr3 = tbl3.rows[0]
        for i, col in enumerate(eri_cols):
            hdr3.cells[i].text = col
            for para in hdr3.cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            tc = hdr3.cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F3F4F6")
            tcPr.append(shd)

        gri_rows = [
            ("GRI 302 (Energy)",      f"Annual solar generation {round(total_kwh/1000, 1)} MWh",       "Reduction in Scope 2 purchased electricity"),
            ("GRI 303 (Water)",       f"Annual rainwater harvest {_fmt_inr(total_rain_kl)} kL",         "Reduction in municipal water consumption"),
            ("GRI 304 (Biodiversity)", f"{_fmt_inr(total_plants)} plants, multiple species",            "Urban biodiversity contribution"),
            ("GRI 305 (Emissions)",   f"{total_co2_t} tCO2e/yr offset",                               "Solar + plantation carbon sequestration"),
        ]
        for gri, est, rel in gri_rows:
            r = tbl3.add_row()
            for i, val in enumerate([gri, est, rel]):
                r.cells[i].text = val
                for para in r.cells[i].paragraphs:
                    for run in para.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

        doc.add_paragraph()
        heading("Verification Pathways", level=2)
        bullet("Indian Carbon Market (ICM): BEE-accredited Carbon Verification Agencies, NCCF")
        bullet("Global VCM: Verra (VCS) or Gold Standard; VVBs: Carbon Check India, TUV Nord")
        bullet("Green Building Ratings: GRIHA (India), IGBC/LEED (international)")

    else:
        heading("6. Community Environmental Impact")
        flats = group.get("colony_flats_count") or (bld_count * 4)
        food_per_family = round(total_food_kg / flats) if flats else 0
        water_per_hh    = round(total_rain_l / flats) if flats else 0

        bullet(f"Households benefited: {_fmt_inr(flats)} families")
        bullet(f"Food security contribution: {_fmt_inr(food_per_family)} kg organic produce/family/year")
        bullet(f"Water savings per household: {_fmt_inr(water_per_hh)} litres/year")
        bullet(f"Energy savings: {_fmt_inr(total_kwh)} units solar/year (common area electricity impact)")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SECTION 7 — METHODOLOGY AND DISCLAIMER
    # ════════════════════════════════════════════════════════════════════
    heading("7. Methodology and Disclaimer")
    body("Calculation methodology: MNRE Solar Radiation Handbook (GHI), IMD long-period rainfall averages, "
         "IS 16190:2014 (biogas standards), CPCB/CGWB Rainwater Harvesting Manual, "
         "ICAR KVK container farming guidelines.", bold=False)
    doc.add_paragraph()
    body(
        "DISCLAIMER: All estimates are based on State 1 Remote Assessment using satellite data and published "
        "Indian regulatory standards. Actual performance will vary based on site conditions, equipment "
        "selection, installation quality, and local DISCOM policies. A formal feasibility study requires "
        "a physical site assessment."
    )

    # ════════════════════════════════════════════════════════════════════
    # SECTION 8 — NEXT STEPS AND CONTACT
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    heading("8. Next Steps and Contact")
    body(
        "Ready to move from estimate to implementation? Sus10 AI connects you with verified vendors "
        "for site assessment and system design."
    )
    doc.add_paragraph()
    body("Contact: gp@sus10.ai  |  sus10.ai", bold=True)

    # ── Write to bytes ───────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
